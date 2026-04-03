from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from models.schemas import AuditReport, SectionAuditResult
import os
import logging

logger = logging.getLogger(__name__)

# 宋体字体名称
FONT_SONGTI = "宋体"


class ReportGenerator:
    """生成审阅报告docx"""

    def _set_run_font(self, run, font_name: str = FONT_SONGTI, size: int = 12,
                      bold: bool = False, italic: bool = False):
        """设置文本run的字体"""
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        # 设置中文字体（需要通过XML设置eastAsia属性）
        r = run._element.rPr
        if r is None:
            from docx.oxml import OxmlElement
            r = OxmlElement('w:rPr')
            run._element.insert(0, r)
        rFonts = r.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            r.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    def _set_heading_font(self, heading, font_name: str = FONT_SONGTI, size: int = 14):
        """设置标题字体"""
        for run in heading.runs:
            self._set_run_font(run, font_name, size=size)

    def _add_paragraph_with_font(self, doc, text: str, bold: bool = False, italic: bool = False):
        """添加带宋体的段落"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        self._set_run_font(run, bold=bold, italic=italic)
        return para

    def generate(self, report: AuditReport, output_filename: str = "审阅报告.docx") -> str:
        """
        生成审阅报告

        Args:
            report: 审阅报告数据
            output_filename: 输出文件名

        Returns:
            生成的文件路径
        """
        doc = Document()

        # 设置文档默认字体为宋体
        style = doc.styles['Normal']
        style.font.name = FONT_SONGTI
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)

        # 标题
        title = doc.add_heading("税务合同审阅报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_heading_font(title, size=24)

        # 整体摘要
        summary_heading = doc.add_heading("审阅摘要", level=1)
        self._set_heading_font(summary_heading, size=18)
        self._add_paragraph_with_font(doc, report.summary)

        # 统计信息
        compliant_count = sum(1 for s in report.sections if s.risk_level == "合规")
        high_risk_count = sum(1 for s in report.sections if s.risk_level == "高风险")
        non_compliant_count = sum(1 for s in report.sections if s.risk_level == "不合规")

        stats_para = doc.add_paragraph()
        run1 = stats_para.add_run(f"合规条款: {compliant_count}  ")
        self._set_run_font(run1)

        run2 = stats_para.add_run(f"高风险条款: {high_risk_count}  ")
        self._set_run_font(run2, bold=True)

        run3 = stats_para.add_run(f"不合规条款: {non_compliant_count}")
        self._set_run_font(run3, bold=True)

        # 各章节详情（仅输出高风险和不合规条款）
        detail_heading = doc.add_heading("详细分析", level=1)
        self._set_heading_font(detail_heading, size=18)

        for section in report.sections:
            # 跳过合规条款，不输出
            if section.risk_level == "合规":
                continue
            self._add_section(doc, section)

        # 保存文件
        output_path = os.path.join(os.getcwd(), output_filename)
        doc.save(output_path)
        logger.info(f"Report saved to {output_path}")

        return output_path

    def _add_section(self, doc: Document, section: SectionAuditResult):
        """添加单个章节的分析结果"""

        # 章节标题
        section_heading = doc.add_heading(section.section_name, level=2)
        self._set_heading_font(section_heading, size=14)

        # 风险等级（颜色标记）
        risk_para = doc.add_paragraph()
        label_run = risk_para.add_run("风险等级: ")
        self._set_run_font(label_run, bold=True)

        risk_run = risk_para.add_run(section.risk_level)
        self._set_run_font(risk_run, bold=True)
        if section.risk_level == "不合规":
            risk_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        elif section.risk_level == "高风险":
            risk_run.font.color.rgb = RGBColor(255, 165, 0)  # 橙色
        else:
            risk_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色

        # 条款原文（使用斜体）
        orig_label_para = doc.add_paragraph()
        orig_label_run = orig_label_para.add_run("条款原文:")
        self._set_run_font(orig_label_run, bold=True)

        orig_content_para = doc.add_paragraph()
        orig_content_run = orig_content_para.add_run(section.original_content)
        self._set_run_font(orig_content_run, italic=True)  # 斜体显示原文

        # 违反法规
        if section.violated_regulations:
            reg_label_para = doc.add_paragraph()
            reg_label_run = reg_label_para.add_run("违反法规:")
            self._set_run_font(reg_label_run, bold=True)

            for reg in section.violated_regulations:
                reg_para = doc.add_paragraph()
                reg_run = reg_para.add_run(f"  • {reg}")
                self._set_run_font(reg_run)

        # 原因分析
        if section.reason:
            reason_label_para = doc.add_paragraph()
            reason_label_run = reason_label_para.add_run("原因分析:")
            self._set_run_font(reason_label_run, bold=True)
            self._add_paragraph_with_font(doc, section.reason)

        # 修改建议
        if section.suggestion:
            suggest_label_para = doc.add_paragraph()
            suggest_label_run = suggest_label_para.add_run("修改建议:")
            self._set_run_font(suggest_label_run, bold=True)

            suggest_para = doc.add_paragraph()
            suggest_run = suggest_para.add_run(section.suggestion)
            self._set_run_font(suggest_run)
            suggest_run.font.color.rgb = RGBColor(0, 100, 0)  # 深绿色

        # 分隔线
        divider_para = doc.add_paragraph()
        divider_run = divider_para.add_run("─" * 37)
        self._set_run_font(divider_run)