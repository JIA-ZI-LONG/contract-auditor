from docx import Document
from models.schemas import ContractSection
import logging

logger = logging.getLogger(__name__)

class ContractParser:
    """解析docx合同文档，提取章节结构"""

    def parse(self, file_path: str) -> list[ContractSection]:
        """
        解析合同文档

        Args:
            file_path: docx文件路径

        Returns:
            章节列表，每个章节包含名称和内容
        """
        doc = Document(file_path)
        sections = []
        current_section_name = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 判断是否为章节标题（Heading样式或以"条款"、"第"开头）
            is_heading = (
                para.style.name.startswith("Heading") or
                text.startswith("第") or
                text.endswith("条款") or
                text.endswith("章") or
                text.endswith("节")
            )

            if is_heading:
                # 保存前一个章节
                if current_section_name and current_content:
                    sections.append(ContractSection(
                        section_name=current_section_name,
                        content="\n".join(current_content)
                    ))
                current_section_name = text
                current_content = []
            else:
                current_content.append(text)

        # 保存最后一个章节
        if current_section_name and current_content:
            sections.append(ContractSection(
                section_name=current_section_name,
                content="\n".join(current_content)
            ))

        # 如果没有章节结构，将全文作为一个章节
        if not sections:
            all_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            if all_text:
                sections.append(ContractSection(
                    section_name="全文",
                    content=all_text
                ))

        logger.info(f"Parsed {len(sections)} sections from contract")
        return sections