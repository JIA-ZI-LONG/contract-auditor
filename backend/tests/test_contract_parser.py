import pytest
from services.contract_parser import ContractParser
from docx import Document
import tempfile
import os


def test_parse_contract_with_sections():
    """测试解析带章节结构的合同"""
    # 创建测试文档
    doc = Document()
    doc.add_heading("付款条款", level=1)
    doc.add_paragraph("买方应在收到发票后30日内支付货款。")
    doc.add_heading("违约条款", level=1)
    doc.add_paragraph("逾期按日利率0.05%计算违约金。")

    # 保存临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    # 解析
    parser = ContractParser()
    sections = parser.parse(temp_path)

    # 清理
    os.unlink(temp_path)

    # 验证
    assert len(sections) == 2
    assert sections[0].section_name == "付款条款"
    assert "发票" in sections[0].content
    assert sections[1].section_name == "违约条款"


def test_parse_contract_without_sections():
    """测试解析无章节的纯文本合同"""
    doc = Document()
    doc.add_paragraph("本合同由甲乙双方签订。")
    doc.add_paragraph("付款方式为银行转账。")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    parser = ContractParser()
    sections = parser.parse(temp_path)
    os.unlink(temp_path)

    # 无章节标题时，合并为"全文"章节
    assert len(sections) == 1
    assert sections[0].section_name == "全文"